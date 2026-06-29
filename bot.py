# -*- coding: utf-8 -*-
import os
import re
import random
import warnings
import subprocess
import time
import traceback
import asyncio
import sys
from datetime import datetime, timedelta, timezone
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup, BotCommand
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ConversationHandler, CallbackQueryHandler
from telegram.warnings import PTBUserWarning
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw, ImageFont
import io

# Настройка UTF-8 для Windows
sys.stdout.reconfigure(encoding='utf-8')

warnings.filterwarnings("ignore", category=PTBUserWarning)

# Состояния диалога
FIO_INPUT, CHANGE_DOCNUM, DOCNUM_INPUT, GENDER, DOC_TYPE, DOC_NUMBER, ADDRESS_CHOICE, ADDRESS = range(8)

# Пути к файлам
TEMPLATE_DOCX = "Доверенность_шаблон.docx"
TEMPLATE_PECHAT = "pechat_shablon.png"

# ===== ФОРМАТИРОВАНИЕ =====

def format_passport(passport_input):
    digits = re.sub(r'\D', '', passport_input)
    if len(digits) == 10:
        return f"паспорт серия {digits[:4]} № {digits[4:]}"
    return passport_input

def format_snils(snils_input):
    digits = re.sub(r'\D', '', snils_input)
    return f"СНИЛС {digits}"

def format_inn(inn_input):
    digits = re.sub(r'\D', '', inn_input)
    return f"ИНН {digits}"

def get_short_name(full_name):
    """Иванов Иван Иванович -> Иванов И.И."""
    parts = full_name.split()
    if len(parts) >= 3:
        return f"{parts[0]} {parts[1][0]}.{parts[2][0]}."
    if len(parts) == 2:
        return f"{parts[0]} {parts[1][0]}."
    return full_name

# ===== ГЕНЕРАЦИЯ =====

def generate_cert():
    prefix = "10711"
    remaining = 39 - len(prefix)
    return prefix + ''.join(str(random.randint(0, 9)) for _ in range(remaining))

def generate_docnum():
    return str(random.randint(100000, 999999))

def generate_signature_image(data):
    if not os.path.exists(TEMPLATE_PECHAT):
        raise FileNotFoundError(f"Нет файла: {TEMPLATE_PECHAT}")
    
    img = Image.open(TEMPLATE_PECHAT)
    draw = ImageDraw.Draw(img)
    
    try:
        font = ImageFont.truetype("times.ttf", 11)
    except:
        try:
            font = ImageFont.truetype("Times New Roman.ttf", 11)
        except:
            font = ImageFont.load_default()
    
    text_color = (72, 64, 176)
    
    draw.text((95, 112), data['cert'], fill=text_color, font=font)
    draw.text((86, 132), data['fio'], fill=text_color, font=font)
    date_text = f"с {data['date_now']} до {data['date_future']}"
    draw.text((109, 151), date_text, fill=text_color, font=font)
    
    img_bytes = io.BytesIO()
    img.save(img_bytes, format='PNG')
    img_bytes.seek(0)
    return img_bytes

# ===== КОНВЕРТАЦИЯ =====

def convert_to_pdf(docx_path):
    pdf_path = docx_path.replace('.docx', '.pdf')
    
    if os.path.exists(pdf_path):
        os.remove(pdf_path)
    
    libreoffice_paths = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    
    libreoffice = None
    for path in libreoffice_paths:
        if os.path.exists(path):
            libreoffice = path
            break
    
    if not libreoffice:
        raise Exception("LibreOffice не найден!")
    
    abs_docx = os.path.abspath(docx_path)
    output_dir = os.path.dirname(abs_docx)
    
    subprocess.Popen(
        [libreoffice, '--headless', '--invisible', '--norestore'],
        stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
    )
    time.sleep(2)
    
    result = subprocess.run(
        [libreoffice, '--headless', '--norestore', '--convert-to', 'pdf', '--outdir', output_dir, abs_docx],
        capture_output=True, text=True, timeout=120
    )
    
    if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 1000:
        return pdf_path
    
    base_name = os.path.splitext(os.path.basename(docx_path))[0]
    expected_pdf = os.path.join(output_dir, base_name + ".pdf")
    
    if os.path.exists(expected_pdf) and os.path.getsize(expected_pdf) > 1000:
        if expected_pdf != pdf_path:
            os.rename(expected_pdf, pdf_path)
        return pdf_path
    
    raise Exception("PDF не создан!")

# ===== ЗАПОЛНЕНИЕ DOCX =====

def fill_docx(data):
    print(f"\n[LOG] === ЗАПОЛНЕНИЕ DOCX ===")
    print(f"[LOG] ФИО: {data['fio']}")
    print(f"[LOG] Номер: {data.get('docnum', 'ИЗ ШАБЛОНА')}")
    
    if not os.path.exists(TEMPLATE_DOCX):
        raise FileNotFoundError(f"Файл не найден: {TEMPLATE_DOCX}")
    
    doc = Document(TEMPLATE_DOCX)
    
    # Колонтитулы
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            if "{{ SIGNATURE }}" in paragraph.text:
                paragraph.text = ""
                signature_img = generate_signature_image(data)
                run = paragraph.add_run()
                run.add_picture(signature_img, width=Inches(3.0))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Основной текст
    for paragraph in doc.paragraphs:
        text = paragraph.text
        
        # Отдельная дата
        if re.match(r'^\s*\d{2}\.\d{2}\.\d{4}\s*$', text):
            paragraph.text = data['date_now']
        
        # Исх. №
        elif "Исх" in text and "№" in text:
            if data.get('docnum'):
                paragraph.text = re.sub(r'(Исх\.?\s*№\s*)\d+', rf'\g<1>{data["docnum"]}', text)
        
        # Документ №
        elif "Документ №" in text:
            if data.get('docnum'):
                paragraph.text = re.sub(r'(Документ\s*№\s*)\d+', rf'\g<1>{data["docnum"]}', text)
        
        # ФИО доверителя — НЕ ТРОГАЕМ БОГДАНОВА, ПОДПИСЬ, ПРЕДСЕДАТЕЛЯ, ЧИСТЮХИНА
        elif (re.search(r'[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+(?:,\s*\d{2}\.\d{2}\.\d{4},\s*(?:женский|мужской))?', text) 
              and "Богданова" not in text 
              and "Первый заместитель" not in text
              and "Председателя" not in text
              and "Чистюхин" not in text):
            print("[LOG] Замена данных доверителя")
            
            old_fio_pattern = r'[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+(?:,\s*\d{2}\.\d{2}\.\d{4},\s*(?:женский|мужской))?'
            
            if data.get('birth') and data.get('gender'):
                replacement = data['fio_birth_gender']
            else:
                replacement = data['fio']
            
            text = re.sub(old_fio_pattern, replacement, text)
            
            # Документ
            if data.get('doc_display'):
                text = re.sub(
                    r',?\s*\[?,\s*паспорт\s+серия\s*\d+\s*№\s*\d+\s*,\s*\]?',
                    f", {data['doc_display']}, ",
                    text
                )
            else:
                text = re.sub(
                    r',?\s*\[?,\s*паспорт\s+серия\s*\d+\s*№\s*\d+\s*,\s*\]?',
                    ", ",
                    text
                )
            
            # Склонение
            if data.get('reg_ending'):
                text = re.sub(r'зарегистрирован[аыо]?\b', data['reg_ending'], text)
            
            # Адрес
            if data.get('address'):
                text = re.sub(
                    r'Краснодарский край,\s*г\.\s*Геленджик,\s*ул\.\s*Леселидзе,\s*д\.\s*\d+,\s*кв\.\s*\d+',
                    data['address'],
                    text
                )
            else:
                text = re.sub(
                    r'\s*по адресу:\s*Краснодарский край,\s*г\.\s*Геленджик,\s*ул\.\s*Леселидзе,\s*д\.\s*\d+,\s*кв\.\s*\d+',
                    '',
                    text
                )
            
            paragraph.text = text
        
        # Даты — НЕ ТРОГАЕМ ПОДПИСЬ
        elif "Чистюхин" not in text and "Первый заместитель" not in text and "Председателя" not in text:
            new_text = re.sub(r'\d{2}\.\d{2}\.2029', data['date_future'], text)
            new_text = re.sub(r'2[1-8]\.\d{2}\.202[6-9]', data['date_now'], new_text)
            if new_text != text:
                paragraph.text = new_text
        
        # Шрифт
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11) if "Документ №" in paragraph.text else Pt(10)
    
    short_name = get_short_name(data['fio']).replace(' ', '_').replace('.', '')
    output_path = f"Доверенность_{short_name}.docx"
    doc.save(output_path)
    return output_path

# ===== TELEGRAM HANDLERS =====

async def help_command(update: Update, context):
    await update.message.reply_text(
        "📋 *Бот для создания доверенностей*\n\n"
        "/start — 🚀 Создать доверенность\n"
        "/cancel — ❌ Отменить\n"
        "/help — ℹ️ Помощь\n\n"
        "*Ввод ФИО:*\n"
        "• Просто текст: `Иванов Иван Иванович`\n"
        "• С датой: `Иванов Иван Иванович, 01.01.1990`\n\n"
        "*Номер доверенности:*\n"
        "• Оставить исходный\n"
        "• Ввести свой",
        parse_mode="Markdown"
    )

async def start(update: Update, context):
    print("\n[LOG] === НОВЫЙ ПОЛЬЗОВАТЕЛЬ ===")
    await update.message.reply_text(
        "👋 Здравствуйте! Я помогу создать доверенность.\n\n"
        "Введите ФИО (можно просто имя или название):\n"
        "Например: Иванов Иван Иванович"
    )
    return FIO_INPUT

async def fio_input(update: Update, context):
    msg = update.message
    text = msg.text.strip()
    await msg.delete()
    
    print(f"[LOG] Получено: {text}")
    
    # Принимаем ЛЮБОЙ текст как ФИО
    context.user_data['fio'] = text
    context.user_data['fio_raw'] = text
    
    # Проверяем, есть ли дата рождения
    parts = text.split(',')
    if len(parts) >= 2 and re.search(r'\d{2}\.\d{2}\.\d{4}', parts[1]):
        context.user_data['fio'] = parts[0].strip()
        context.user_data['birth'] = parts[1].strip()
        context.user_data['fio_raw'] = f"{parts[0].strip()}, {parts[1].strip()}"
        has_birth = True
    else:
        context.user_data['birth'] = ''
        has_birth = False
    
    print(f"[LOG] ФИО: {context.user_data['fio']}")
    print(f"[LOG] Дата рождения: {context.user_data.get('birth', 'НЕТ')}")
    
    # Кнопки: изменить номер или оставить
    keyboard = [
        [InlineKeyboardButton("📝 Ввести свой номер доверенности", callback_data="change_docnum")],
        [InlineKeyboardButton("✅ Оставить номер из шаблона", callback_data="keep_docnum")],
    ]
    
    await update.message.reply_text(
        f"✅ ФИО: {context.user_data['fio']}\n"
        + (f"📅 Дата рождения: {context.user_data['birth']}\n" if has_birth else "")
        + f"\n📋 Номер доверенности:",
        reply_markup=InlineKeyboardMarkup(keyboard)
    )
    return CHANGE_DOCNUM

async def change_docnum(update: Update, context):
    query = update.callback_query
    await query.answer()
    
    choice = query.data
    
    if choice == "keep_docnum":
        context.user_data['docnum'] = ''
        
        if context.user_data.get('birth'):
            keyboard = [
                [InlineKeyboardButton("👨 Мужской", callback_data="male"),
                 InlineKeyboardButton("👩 Женский", callback_data="female")]
            ]
            await query.edit_message_text(
                f"✅ ФИО: {context.user_data['fio']}\n"
                f"📅 Дата рождения: {context.user_data['birth']}\n"
                f"📋 Номер: ИЗ ШАБЛОНА\n\n"
                f"Выберите пол:",
                reply_markup=InlineKeyboardMarkup(keyboard)
            )
            return GENDER
        else:
            return await go_to_docs(update, context)
    
    await query.edit_message_text(
        f"✅ ФИО: {context.user_data['fio']}\n\n"
        f"📝 Введите номер доверенности (6 цифр):\n"
        f"Например: 123456"
    )
    return DOCNUM_INPUT

async def docnum_input(update: Update, context):
    msg = update.message
    docnum = msg.text.strip()
    await msg.delete()
    
    docnum = re.sub(r'\D', '', docnum)
    if not docnum:
        docnum = generate_docnum()
    
    context.user_data['docnum'] = docnum
    
    print(f"[LOG] Номер: {docnum}")
    
    if context.user_data.get('birth'):
        keyboard = [
            [InlineKeyboardButton("👨 Мужской", callback_data="male"),
             InlineKeyboardButton("👩 Женский", callback_data="female")]
        ]
        await update.message.reply_text(
            f"✅ ФИО: {context.user_data['fio']}\n"
            f"📅 Дата рождения: {context.user_data['birth']}\n"
            f"📋 Номер: {docnum}\n\n"
            f"Выберите пол:",
            reply_markup=InlineKeyboardMarkup(keyboard)
        )
        return GENDER
    else:
        return await go_to_docs(update, context)

async def go_to_docs(update, context):
    keyboard = [
        [InlineKeyboardButton("📘 Паспорт РФ", callback_data="passport")],
        [InlineKeyboardButton("📋 СНИЛС", callback_data="snils")],
        [InlineKeyboardButton("📄 ИНН", callback_data="inn")],
        [InlineKeyboardButton("🚫 Без документов", callback_data="none")],
    ]
    
    text = f"✅ ФИО: {context.user_data['fio']}\n"
    if context.user_data.get('docnum'):
        text += f"📋 Номер: {context.user_data['docnum']}\n"
    else:
        text += f"📋 Номер: ИЗ ШАБЛОНА\n"
    text += "\n📝 Выберите тип документа:"
    
    if update.callback_query:
        await update.callback_query.edit_message_text(text, reply_markup=InlineKeyboardMarkup(keyboard))
    else:
        await update.message.reply_text(text, reply_markup=InlineKeyboardMarkup(keyboard))
    
    return DOC_TYPE

async def gender(update: Update, context):
    query = update.callback_query
    await query.answer()
    
    if query.data == "male":
        context.user_data['gender'] = "мужской"
        context.user_data['reg_ending'] = "зарегистрирован"
        emoji = "👨"
    else:
        context.user_data['gender'] = "женский"
        context.user_data['reg_ending'] = "зарегистрирована"
        emoji = "👩"
    
    context.user_data['fio_birth_gender'] = f"{context.user_data['fio_raw']}, {context.user_data['gender']}"
    
    print(f"[LOG] Пол: {context.user_data['gender']}")
    
    return await go_to_docs(update, context)

async def doc_type(update: Update, context):
    query = update.callback_query
    await query.answer()
    
    doc_key = query.data
    
    if doc_key == "none":
        context.user_data['doc_display'] = ''
        context.user_data['doc_type'] = 'none'
    else:
        doc_type_map = {
            "passport": ("паспорт", "📘 Паспорт РФ", "10 цифр"),
            "snils": ("СНИЛС", "📋 СНИЛС", "11 цифр"),
            "inn": ("ИНН", "📄 ИНН", "12 цифр"),
        }
        doc_name, doc_emoji, doc_hint = doc_type_map[doc_key]
        context.user_data['doc_type'] = doc_key
        
        await query.edit_message_text(
            f"✅ ФИО: {context.user_data['fio']}\n"
            f"{doc_emoji} Тип документа: {doc_name}\n\n"
            f"📝 Введите {doc_name} ({doc_hint}):"
        )
        return DOC_NUMBER
    
    keyboard = [
        [InlineKeyboardButton("📍 Ввести адрес", callback_data="with_address")],
        [InlineKeyboardButton("🚫 Без адреса", callback_data="no_address")],
    ]
    await query.edit_message_text(
        f"✅ ФИО: {context.user_data['fio']}\n"
        f"🚫 Без документов\n\n"
        f"📍 Выберите вариант адреса:",
        reply_markup=InlineKeyboardMarkup(keyboard)
    )
    return ADDRESS_CHOICE

async def doc_number(update: Update, context):
    msg = update.message
    doc_input = msg.text.strip()
    await msg.delete()
    
    doc_type = context.user_data['doc_type']
    
    if doc_type == "passport":
        formatted = format_passport(doc_input)
    elif doc_type == "snils":
        formatted = format_snils(doc_input)
    elif doc_type == "inn":
        formatted = format_inn(doc_input)
    else:
        formatted = doc_input
    
    context.user_data['doc_display'] = formatted
    
    keyboard = [
        [InlineKeyboardButton("📍 Ввести адрес", callback_data="with_address")],
        [InlineKeyboardButton("🚫 Без адреса", callback_data="no_address")],
    ]
    
    await update.message.reply_text(
        f"✅ Документ: {formatted}\n\n"
        f"📍 Выберите вариант адреса:",
        reply_markup=InlineKeyboardMarkup(keyboard)
    )
    return ADDRESS_CHOICE

async def address_choice(update: Update, context):
    query = update.callback_query
    await query.answer()
    
    choice = query.data
    
    if choice == "no_address":
        context.user_data['address'] = ''
        await query.edit_message_text("⏳ Генерирую доверенность...")
        return await generate_and_send(update, context)
    
    await query.edit_message_text(
        f"✅ Документ: {context.user_data.get('doc_display', 'нет')}\n\n"
        f"📍 Введите адрес регистрации:"
    )
    return ADDRESS

async def address(update: Update, context):
    msg = update.message
    context.user_data['address'] = msg.text.strip()
    await msg.delete()
    
    await update.message.reply_text("⏳ Генерирую доверенность...")
    return await generate_and_send(update, context)

async def generate_and_send(update: Update, context):
    print(f"[LOG] Генерация...")
    
    moscow_tz = timezone(timedelta(hours=3))
    today = datetime.now(moscow_tz)
    future = today.replace(year=today.year + 3)
    
    data = {
        'fio': context.user_data['fio'],
        'fio_raw': context.user_data.get('fio_raw', context.user_data['fio']),
        'birth': context.user_data.get('birth', ''),
        'gender': context.user_data.get('gender', ''),
        'reg_ending': context.user_data.get('reg_ending', 'зарегистрирован'),
        'fio_birth_gender': context.user_data.get('fio_birth_gender', context.user_data['fio']),
        'doc_display': context.user_data.get('doc_display', ''),
        'docnum': context.user_data.get('docnum', ''),
        'address': context.user_data.get('address', ''),
        'date_now': today.strftime("%d.%m.%Y"),
        'date_future': future.strftime("%d.%m.%Y"),
        'cert': generate_cert(),
    }
    
    try:
        docx_path = fill_docx(data)
        
        if update.callback_query:
            status = await update.callback_query.edit_message_text("🔄 Конвертирую в PDF...")
        else:
            status = await update.message.reply_text("🔄 Конвертирую в PDF...")
        
        pdf_path = convert_to_pdf(docx_path)
        
        short_name = get_short_name(data['fio'])
        send_filename = f"Доверенность {short_name}.pdf"
        
        with open(pdf_path, 'rb') as f:
            await update.effective_chat.send_document(
                document=f,
                filename=send_filename,
                caption="✅ Доверенность успешно создана!"
            )
        
        if update.callback_query:
            await update.callback_query.delete_message()
        else:
            await status.delete()
        
        if os.path.exists(docx_path): os.remove(docx_path)
        if os.path.exists(pdf_path): os.remove(pdf_path)
        
    except Exception as e:
        print(f"[ERROR] {e}")
        traceback.print_exc()
        if update.callback_query:
            await update.callback_query.edit_message_text(f"❌ Ошибка: {e}")
        else:
            await update.message.reply_text(f"❌ Ошибка: {e}")
    
    return ConversationHandler.END

async def cancel(update: Update, context):
    await update.message.reply_text("❌ Отменено. /start")
    return ConversationHandler.END

async def set_bot_commands(app):
    commands = [
        BotCommand("start", "🚀 Создать доверенность"),
        BotCommand("cancel", "❌ Отменить"),
        BotCommand("help", "ℹ️ Помощь"),
    ]
    await app.bot.set_my_commands(commands)

def main():
    TOKEN = "8696458590:AAGySeTkuFMgAMBGJE9Rqi8kaqeJwuy_H_4"
    
    print("\n" + "="*50)
    print("БОТ ДОВЕРЕННОСТЕЙ")
    print("="*50)
    print(f"Шаблон: {os.path.exists(TEMPLATE_DOCX)}")
    print(f"Печать: {os.path.exists(TEMPLATE_PECHAT)}")
    print("="*50 + "\n")
    
    app = Application.builder().token(TOKEN).build()
    app.add_handler(CommandHandler("help", help_command))
    
    conv = ConversationHandler(
        entry_points=[CommandHandler('start', start)],
        states={
            FIO_INPUT: [MessageHandler(filters.TEXT & ~filters.COMMAND, fio_input)],
            CHANGE_DOCNUM: [CallbackQueryHandler(change_docnum)],
            DOCNUM_INPUT: [MessageHandler(filters.TEXT & ~filters.COMMAND, docnum_input)],
            GENDER: [CallbackQueryHandler(gender)],
            DOC_TYPE: [CallbackQueryHandler(doc_type)],
            DOC_NUMBER: [MessageHandler(filters.TEXT & ~filters.COMMAND, doc_number)],
            ADDRESS_CHOICE: [CallbackQueryHandler(address_choice)],
            ADDRESS: [MessageHandler(filters.TEXT & ~filters.COMMAND, address)],
        },
        fallbacks=[CommandHandler('cancel', cancel)],
    )
    app.add_handler(conv)
    
    async def setup():
        await set_bot_commands(app)
    asyncio.get_event_loop().run_until_complete(setup())
    
    print("Бот запущен!\n")
    app.run_polling()

if __name__ == '__main__':
    main()